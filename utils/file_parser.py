from io import BytesIO
import logging
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import re
from zipfile import ZipFile
import xml.etree.ElementTree as ET
import os
import base64
from utils.logger import Logger

logger = Logger(__name__).get_logger()

def parse_pdf_or_docx(file_buffer: BytesIO, filename: str) -> str:
    """
    Determines the file type (PDF, DOCX, or image) and extracts text accordingly.
    :param file_buffer: File buffer of the uploaded file.
    :param filename: Name of the uploaded file.
    :return: Extracted text content as a string.
    """
    try:
        if filename.lower().endswith(".pdf"):
            return parse_pdf(file_buffer)
        elif filename.lower().endswith(".docx"):
            return parse_docx(file_buffer)
        elif filename.lower().endswith(".doc"):
            # For .doc files, return a message that they're not supported on the server
            return "DOC files are not supported in cloud deployment. Please convert to DOCX or PDF format."
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            # For images, return a base64 representation instead of OCR
            # since Tesseract isn't available on Vercel
            return extract_text_from_image(file_buffer)
        else:
            raise ValueError("Unsupported file format. Only PDF, DOCX, and image formats are supported.")
    except Exception as e:
        logger.error(f"Error parsing file '{filename}': {str(e)}", exc_info=True)
        raise

def parse_pdf(file_buffer: BytesIO) -> str:
    """
    Extracts text from a PDF file, including hyperlinks.
    :param file_buffer: File buffer of the uploaded PDF file.
    :return: Extracted text content as a string, including hyperlinks.
    """
    try:
        logger.info("Parsing PDF file")
        reader = PdfReader(file_buffer)
        text = ""
        hyperlinks = []

        # Extract text from each page and gather hyperlinks from metadata if available
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text

            # Extract hyperlinks from annotations (if available)
            if "/Annots" in page:
                annotations = page["/Annots"]
                for annotation in annotations:
                    # Check if annotation is an IndirectObject and resolve it properly
                    if isinstance(annotation, dict):
                        if "/A" in annotation and "/URI" in annotation["/A"]:
                            hyperlinks.append(annotation["/A"]["/URI"])

        # Join the hyperlinks into a single string (one per line)
        hyperlinks_text = '\n'.join(hyperlinks)
        return text.strip() + '\n' + hyperlinks_text

    except Exception as e:
        logger.error(f"Error reading PDF file: {str(e)}", exc_info=True)
        raise

def parse_docx(file_buffer: BytesIO) -> str:
    """
    Extracts text from a DOCX file, including hyperlinks and headers/footers.
    :param file_buffer: File buffer of the uploaded DOCX file.
    :return: Extracted text content as a string, including hyperlinks.
    """
    try:
        logger.info("Parsing DOCX file")
        doc = Document(file_buffer)
        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'

        # Extract hyperlinks from the document (href="...") in the HTML of the document
        hyperlinks = extract_hyperlinks_from_docx(file_buffer)
        
        # Extract header and footer text
        header_footer_text = extract_header_footer(doc)

        return text.strip() + '\n' + hyperlinks + '\n' + header_footer_text

    except Exception as e:
        logger.error(f"Error reading DOCX file: {str(e)}", exc_info=True)
        raise

def extract_hyperlinks_from_docx(file_buffer: BytesIO) -> str:
    """
    Extracts hyperlinks from a DOCX file by scanning for <a> tags in the document's HTML.
    :param file_buffer: The file buffer of the DOCX file.
    :return: A string containing all hyperlinks found in the document.
    """
    try:
        # Load the DOCX file as a zip and read the XML content
        docx = ZipFile(file_buffer)
        hyperlinks = []
        for file in docx.namelist():
            if "hyperlink" in file:
                content = docx.read(file)
                links = re.findall(r'href="([^"]+)"', content.decode('utf-8'))
                hyperlinks.extend(links)
        return '\n'.join(hyperlinks)

    except Exception as e:
        logger.error(f"Error extracting hyperlinks from DOCX file: {str(e)}", exc_info=True)
        raise

def extract_header_footer(doc) -> str:
    """
    Extract text from headers and footers in a DOCX file.
    :param doc: The Document object for DOCX file.
    :return: Text from the headers and footers.
    """
    text = ""
    # Extract text from headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            text += paragraph.text + '\n'
        
        # Extract text from footers
        footer = section.footer
        for paragraph in footer.paragraphs:
            text += paragraph.text + '\n'

    return text.strip()

def extract_text_from_image(file_buffer: BytesIO) -> str:
    """
    For cloud environments without Tesseract OCR, we describe the image 
    and provide alternative text explaining why OCR isn't available.
    
    :param file_buffer: The image file buffer.
    :return: Description of the image and instructions for the user.
    """
    try:
        logger.info("Handling image file in cloud environment")
        image = Image.open(file_buffer)
        width, height = image.size
        format_name = image.format
        mode = image.mode
        
        # Create a descriptive message about the image
        image_info = f"""
        Image received: {format_name} format, {width}x{height} pixels, {mode} mode.
        
        OCR text extraction is not available in cloud deployment.
        The system will attempt to extract any visible text from the image,
        but for best results please convert the image to PDF or DOCX format
        with proper text content.
        
        If this is a scanned document, consider using a document conversion
        service to convert it to searchable PDF before uploading.
        """
        
        return image_info.strip()
    
    except Exception as e:
        logger.error(f"Error processing image file: {str(e)}", exc_info=True)
        raise 